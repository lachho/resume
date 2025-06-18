"""Enhanced resume file parsers with ATS compatibility analysis."""

import logging
from pathlib import Path
from typing import Optional, Dict, Any, Tuple

logger = logging.getLogger(__name__)

class ParseResult:
    """Class to hold parsing results with metadata."""
    def __init__(self, text: str, metadata: Dict[str, Any]):
        self.text = text
        self.metadata = metadata

def parse_pdf_enhanced(file_path: Path) -> Optional[ParseResult]:
    """Parse PDF file with enhanced analysis for ATS compatibility."""
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(file_path)
        text_content = []
        metadata = {
            'has_images': False,
            'image_count': 0,
            'has_columns': False,
            'page_count': len(doc),
            'parsing_issues': []
        }
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            
            # Extract text
            text = page.get_text()
            text_content.append(text)
            
            # Check for images
            images = page.get_images()
            if images:
                metadata['has_images'] = True
                metadata['image_count'] += len(images)
                logger.debug(f"Found {len(images)} images on page {page_num + 1}")
            
            # Check for columns (heuristic: text blocks with significant horizontal separation)
            blocks = page.get_text("dict")["blocks"]
            text_blocks = [block for block in blocks if "lines" in block]
            
            if len(text_blocks) >= 2:
                # Check if text blocks are side by side (potential columns)
                x_positions = [block["bbox"][0] for block in text_blocks]
                y_positions = [block["bbox"][1] for block in text_blocks]
                
                # Simple column detection: multiple text blocks at similar y-levels but different x-levels
                for i, block1 in enumerate(text_blocks):
                    for j, block2 in enumerate(text_blocks[i+1:], i+1):
                        y_overlap = min(block1["bbox"][3], block2["bbox"][3]) - max(block1["bbox"][1], block2["bbox"][1])
                        x_separation = abs(block1["bbox"][0] - block2["bbox"][0])
                        
                        # If blocks overlap vertically but are far apart horizontally, likely columns
                        if y_overlap > 50 and x_separation > 100:
                            metadata['has_columns'] = True
                            logger.debug(f"Detected potential columns on page {page_num + 1}")
                            break
                    if metadata['has_columns']:
                        break
        
        doc.close()
        full_text = "\n".join(text_content)
        
        # Additional parsing quality checks
        if len(full_text.strip()) < 50:
            metadata['parsing_issues'].append("Very little text extracted - possible parsing problems")
        
        # Check for excessive special characters (often indicates formatting issues)
        special_char_ratio = sum(1 for char in full_text if not char.isalnum() and not char.isspace()) / len(full_text) if full_text else 0
        if special_char_ratio > 0.3:
            metadata['parsing_issues'].append("High ratio of special characters detected")
        
        logger.debug(f"PDF parsing complete: {len(full_text)} chars, {metadata['image_count']} images, columns: {metadata['has_columns']}")
        return ParseResult(full_text, metadata)
        
    except ImportError:
        logger.error("PyMuPDF (fitz) not installed. Please install it to parse PDF files.")
        return None
    except Exception as e:
        logger.error(f"Error parsing PDF file: {str(e)}")
        return None

def parse_docx_enhanced(file_path: Path) -> Optional[ParseResult]:
    """Parse DOCX file with enhanced analysis for ATS compatibility."""
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document(file_path)
        text_content = []
        metadata = {
            'has_images': False,
            'image_count': 0,
            'has_columns': False,
            'has_tables': False,
            'table_count': 0,
            'parsing_issues': []
        }
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Check for and extract text from tables
        tables = doc.tables
        if tables:
            metadata['has_tables'] = True
            metadata['table_count'] = len(tables)
            
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
        
        # Check for images in document
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                metadata['has_images'] = True
                metadata['image_count'] += 1
        
        # Check for columns (look at section formatting)
        for section in doc.sections:
            if hasattr(section, '_sectPr') and section._sectPr.xpath('.//w:cols'):
                cols_element = section._sectPr.xpath('.//w:cols')[0]
                if cols_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num'):
                    num_cols = int(cols_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num', '1'))
                    if num_cols > 1:
                        metadata['has_columns'] = True
                        logger.debug(f"Document has {num_cols} columns")
                        break
        
        full_text = "\n".join(text_content)
        
        # Additional parsing quality checks
        if len(full_text.strip()) < 50:
            metadata['parsing_issues'].append("Very little text extracted")
        
        logger.debug(f"DOCX parsing complete: {len(full_text)} chars, {metadata['image_count']} images, columns: {metadata['has_columns']}")
        return ParseResult(full_text, metadata)
        
    except ImportError:
        logger.error("python-docx not installed. Please install it to parse DOCX files.")
        return None
    except Exception as e:
        logger.error(f"Error parsing DOCX file: {str(e)}")
        return None

def parse_txt_enhanced(file_path: Path) -> Optional[ParseResult]:
    """Parse text file with basic analysis."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        metadata = {
            'has_images': False,
            'image_count': 0,
            'has_columns': False,
            'parsing_issues': []
        }
        
        # Simple column detection for text files (heuristic: many short lines)
        lines = text.split('\n')
        short_lines = sum(1 for line in lines if 5 < len(line.strip()) < 40)
        if len(lines) > 10 and short_lines / len(lines) > 0.6:
            metadata['has_columns'] = True
            metadata['parsing_issues'].append("Text appears to be formatted in columns")
        
        logger.debug(f"Text file parsing complete: {len(text)} chars")
        return ParseResult(text, metadata)
        
    except Exception as e:
        logger.error(f"Error reading text file: {str(e)}")
        return None

def parse_resume_file(file_path: Path) -> Optional[str]:
    """Legacy function for backward compatibility - returns only text."""
    result = parse_resume_file_enhanced(file_path)
    return result.text if result else None

def parse_resume_file_enhanced(file_path: Path) -> Optional[ParseResult]:
    """Parse resume file with enhanced analysis based on file extension."""
    if not file_path.exists():
        logger.error(f"File does not exist: {file_path}")
        return None
    
    extension = file_path.suffix.lower()
    
    if extension == '.pdf':
        return parse_pdf_enhanced(file_path)
    elif extension == '.docx':
        return parse_docx_enhanced(file_path)
    elif extension == '.txt':
        return parse_txt_enhanced(file_path)
    else:
        logger.error(f"Unsupported file format: {extension}")
        return None 